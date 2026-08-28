# -*- coding: utf-8 -*-
"""Módulo para la generación de borradores oficiales de cartas en formato Microsoft Word (.docx).

Diseñado exactamente según el formato oficial y membrete de CHINA GEZHOUBA GROUP COMPANY LIMITED
para la obra 'Hospital Leoncio Prado de Huamachuco'.
"""
import io
import os
import re
from datetime import datetime, date
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

MESES = [
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "setiembre", "octubre", "noviembre", "diciembre"
]


def _formatear_fecha_larga(fecha_val) -> str:
    """Convierte fecha a formato formal 'Huamachuco, 28 de agosto del 2026'."""
    if not fecha_val:
        d = date.today()
    elif isinstance(fecha_val, (datetime, date)):
        d = fecha_val
    else:
        try:
            s = str(fecha_val).strip().split("T")[0]
            parts = [int(p) for p in s.split("-")]
            d = date(parts[0], parts[1], parts[2])
        except Exception:
            d = date.today()

    mes_nombre = MESES[d.month - 1]
    return f"Huamachuco, {d.day} de {mes_nombre} del {d.year}"


def _set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    """Ajusta márgenes internos de una celda de tabla en dxa."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def _set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    """Aplica bordes delgados y elegantes a la tabla de metadatos."""
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)


def _eliminar_bordes_tabla(table):
    """Elimina todos los bordes visibles de una tabla."""
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="none"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'<w:insideH w:val="none"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)


def _sanitize_xml_str(val) -> str:
    """Elimina caracteres de control ASCII (excepto \\t, \\n, \\r) que corrompen el XML de Word."""
    if val is None:
        return ""
    s = str(val)
    return re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', s)


def generar_carta_docx(payload: dict, config: dict | None = None) -> io.BytesIO:
    """Genera un archivo Word (.docx) con el formato oficial idéntico a la plantilla de obra CGGC.

    Args:
        payload: Datos de la carta (n_documento, fecha, emisor, receptor, dirigido_a,
                 asunto, referencia, especialidad, observacion, etc.).
        config: Configuración institucional (nombre de contratista, obra, etc.).
    """
    config = config or {}
    payload = payload or {}

    contratista = _sanitize_xml_str(config.get("empresa_nombre") or "CHINA GEZHOUBA GROUP COMPANY LIMITED SUCURSAL PERÚ")
    nombre_obra = _sanitize_xml_str(
        config.get("project_title_full") or
        "“MEJORAMIENTO Y AMPLIACIÓN DE LOS SERVICIOS DE SALUD DEL HOSPITAL DE APOYO LEONCIO PRADO, "
        "DISTRITO DE HUAMACHUCO, PROVINCIA SÁNCHEZ CARRIÓN – LA LIBERTAD” con CUI 2335905"
    )

    doc = docx.Document()

    # 1. Configuración de página y márgenes (A4, márgenes exactos de la muestra oficial)
    section = doc.sections[0]
    section.page_width = Inches(8.27)   # A4 Width
    section.page_height = Inches(11.69) # A4 Height
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.8)

    # Configuración de estilos base (Calibri 10.5 pt, color texto formal #111111)
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x11, 0x11, 0x11)

    # ----------------------------------------------------
    # 1. MEMBRETE OFICIAL (BANNER CON LOGOS CEEC Y CGGC)
    # ----------------------------------------------------
    import base64
    from PIL import Image

    custom_logo_word = config.get("logo_membrete_word") or payload.get("logo_membrete_word")
    banner_stream = None

    if custom_logo_word and str(custom_logo_word).strip():
        raw_val = str(custom_logo_word).strip()
        if raw_val.startswith("data:image"):
            try:
                b64_data = raw_val.split(",", 1)[1]
                decoded = base64.b64decode(b64_data)
                # Validar con PIL que sea una imagen real no corrupta
                img_test = Image.open(io.BytesIO(decoded))
                img_test.verify()
                banner_stream = io.BytesIO(decoded)
            except Exception:
                banner_stream = None
        elif os.path.exists(raw_val):
            try:
                img_test = Image.open(raw_val)
                img_test.verify()
                banner_stream = raw_val
            except Exception:
                banner_stream = None

    if not banner_stream:
        banner_paths = [
            os.path.join(os.path.dirname(__file__), "cggc_banner.png"),
            "cggc_banner.png",
            "/app/cggc_banner.png"
        ]
        for bp in banner_paths:
            if os.path.exists(bp):
                try:
                    img_test = Image.open(bp)
                    img_test.verify()
                    banner_stream = bp
                    break
                except Exception:
                    continue

    p_banner = doc.add_paragraph()
    p_banner.paragraph_format.space_before = Pt(0)
    p_banner.paragraph_format.space_after = Pt(4)

    image_placed = False
    if banner_stream:
        try:
            if isinstance(banner_stream, io.BytesIO):
                banner_stream.seek(0)
            r_img = p_banner.add_run()
            r_img.add_picture(banner_stream, width=Inches(6.4))
            image_placed = True
        except Exception:
            image_placed = False

    if not image_placed:
        # Fallback tipográfico idéntico si la imagen no estuviera disponible
        tbl_hdr = doc.add_table(rows=1, cols=2)
        _eliminar_bordes_tabla(tbl_hdr)
        tbl_hdr.autofit = False
        tbl_hdr.cell(0, 0).width = Inches(5.0)
        tbl_hdr.cell(0, 1).width = Inches(1.4)
        p_c1 = tbl_hdr.cell(0, 0).paragraphs[0]
        r_c1_cn = p_c1.add_run("中国葛洲坝集团股份有限公司秘鲁分公司\n")
        r_c1_cn.font.name = 'Arial'
        r_c1_cn.font.size = Pt(13)
        r_c1_cn.font.bold = True
        r_c1_es = p_c1.add_run("CHINA GEZHOUBA GROUP COMPANY LIMITED SUCURSAL PERÚ")
        r_c1_es.font.name = 'Arial'
        r_c1_es.font.size = Pt(9.5)
        r_c1_es.font.bold = True
        r_c1_es.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    # ----------------------------------------------------
    # 2. FECHA Y LUGAR (Alineado a la derecha)
    # ----------------------------------------------------
    p_fecha = doc.add_paragraph()
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_fecha.paragraph_format.space_before = Pt(8)
    p_fecha.paragraph_format.space_after = Pt(14)
    r_fecha = p_fecha.add_run(_formatear_fecha_larga(payload.get("fecha")))
    r_fecha.font.size = Pt(10.5)

    # ----------------------------------------------------
    # 3. IDENTIFICADOR Y DESTINATARIO
    # ----------------------------------------------------
    num_doc = _sanitize_xml_str(payload.get("n_documento") or "Carta N°378-2026-CGGC-HLP-RO").strip()
    if not num_doc.lower().startswith("carta"):
        num_doc = f"Carta {num_doc}"

    p_doc = doc.add_paragraph()
    p_doc.paragraph_format.space_before = Pt(0)
    p_doc.paragraph_format.space_after = Pt(4)
    r_doc = p_doc.add_run(num_doc)
    r_doc.font.bold = True
    r_doc.font.size = Pt(10.5)

    p_dest = doc.add_paragraph()
    p_dest.paragraph_format.line_spacing = 1.15
    p_dest.paragraph_format.space_before = Pt(0)
    p_dest.paragraph_format.space_after = Pt(10)

    r_sres = p_dest.add_run("Señores\n")
    
    dest_input = _sanitize_xml_str(payload.get("dirigido_a") or payload.get("destinatario") or "SUPERVISIÓN").upper()
    if "SUPERVIS" in dest_input or "CONSULTOR" in dest_input or "CARRI" in dest_input:
        nombre_dest = "CONSORCIO CONSULTOR CARRIÓN"
        atencion_persona = "Ing. Javier Julio Quispe Gonzales"
        atencion_cargo = "Jefe de Supervisión"
    elif "PRONIS" in dest_input:
        nombre_dest = "PROGRAMA NACIONAL DE INVERSIONES EN SALUD - PRONIS"
        atencion_persona = "Ing. [Coordinador de Obra PRONIS]"
        atencion_cargo = "Coordinador de Obra"
    elif "MUNICIPAL" in dest_input or "MPSC" in dest_input:
        nombre_dest = "MUNICIPALIDAD PROVINCIAL SÁNCHEZ CARRIÓN"
        atencion_persona = "Ing. [Gerente de Desarrollo Urbano]"
        atencion_cargo = "Gerente de Desarrollo Urbano"
    else:
        nombre_dest = dest_input
        atencion_persona = "Ing. [Responsable / Supervisor]"
        atencion_cargo = "Atención Oficial"

    r_ent = p_dest.add_run(f"{_sanitize_xml_str(nombre_dest)}\n")
    r_ent.font.bold = True
    r_dist = p_dest.add_run("Distrito de Huamachuco\n")
    r_pres = p_dest.add_run("Presente. -")
    r_pres.font.bold = True
    r_pres.font.underline = True

    # ----------------------------------------------------
    # 4. CUADRO DE METADATOS (TABLA CON BORDES GRISES DELGADOS)
    # ----------------------------------------------------
    asunto_txt = _sanitize_xml_str(payload.get("asunto") or "LEVANTAMIENTO DE OBSERVACIONES TÉCNICAS.").strip().upper()
    esp_txt = _sanitize_xml_str(payload.get("especialidad") or payload.get("especialidad_norm") or "GENERAL")
    ref_input = _sanitize_xml_str(payload.get("referencia") or payload.get("referencias") or "")

    table_box = doc.add_table(rows=4, cols=2)
    table_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table_box, color="BFBFBF", sz="4", val="single")
    table_box.autofit = False

    w_col0 = Inches(1.2)
    w_col1 = Inches(5.27)

    for row in table_box.rows:
        row.cells[0].width = w_col0
        row.cells[1].width = w_col1
        _set_cell_margins(row.cells[0], top=80, bottom=80, left=100, right=80)
        _set_cell_margins(row.cells[1], top=80, bottom=80, left=80, right=100)

    # Fila 1: Atención
    p_at_l = table_box.cell(0, 0).paragraphs[0]
    p_at_l.paragraph_format.space_before = Pt(0)
    p_at_l.paragraph_format.space_after = Pt(0)
    r_at_l = p_at_l.add_run("Atención")
    r_at_l.font.bold = True

    p_at_r = table_box.cell(0, 1).paragraphs[0]
    p_at_r.paragraph_format.line_spacing = 1.15
    p_at_r.paragraph_format.space_before = Pt(0)
    p_at_r.paragraph_format.space_after = Pt(0)
    p_at_r.add_run(f":  {_sanitize_xml_str(atencion_persona)}\n   {_sanitize_xml_str(atencion_cargo)}")

    # Fila 2: Asunto
    p_as_l = table_box.cell(1, 0).paragraphs[0]
    p_as_l.paragraph_format.space_before = Pt(0)
    p_as_l.paragraph_format.space_after = Pt(0)
    r_as_l = p_as_l.add_run("Asunto")
    r_as_l.font.bold = True

    p_as_r = table_box.cell(1, 1).paragraphs[0]
    p_as_r.paragraph_format.space_before = Pt(0)
    p_as_r.paragraph_format.space_after = Pt(0)
    p_as_r.add_run(":  ")
    r_as_txt = p_as_r.add_run(asunto_txt)
    r_as_txt.font.bold = True

    # Fila 3: OBRA
    p_ob_l = table_box.cell(2, 0).paragraphs[0]
    p_ob_l.paragraph_format.space_before = Pt(0)
    p_ob_l.paragraph_format.space_after = Pt(0)
    r_ob_l = p_ob_l.add_run("OBRA")
    r_ob_l.font.bold = True

    p_ob_r = table_box.cell(2, 1).paragraphs[0]
    p_ob_r.paragraph_format.space_before = Pt(0)
    p_ob_r.paragraph_format.space_after = Pt(0)
    p_ob_r.paragraph_format.line_spacing = 1.1
    p_ob_r.add_run(f":  {nombre_obra}")

    # Fila 4: Referencia
    p_rf_l = table_box.cell(3, 0).paragraphs[0]
    p_rf_l.paragraph_format.space_before = Pt(0)
    p_rf_l.paragraph_format.space_after = Pt(0)
    r_rf_l = p_rf_l.add_run("Referencia")
    r_rf_l.font.bold = True

    p_rf_r = table_box.cell(3, 1).paragraphs[0]
    p_rf_r.paragraph_format.line_spacing = 1.15
    p_rf_r.paragraph_format.space_before = Pt(0)
    p_rf_r.paragraph_format.space_after = Pt(0)

    # Construir lista de referencias formales con incisos a), b), c)...
    refs_list = ["Contrato de Ejecución de Obra N°003-2024-PRONIS"]
    if ref_input and ref_input.strip() and ref_input != "—":
        # Si vienen múltiples separadas por comas o punto y coma
        for r_part in re.split(r"[,;\n]+", ref_input):
            rp = _sanitize_xml_str(r_part).strip()
            if rp and rp not in refs_list:
                refs_list.append(rp)
    else:
        refs_list.append("Carta N°432-2026-CCC/JQG.")

    # Agregar referencia al informe interno del especialista
    esp_sigla = re.sub(r"[^A-Za-z0-9]", "", esp_txt)[:4].upper() or "GEN"
    refs_list.append(f"Informe N°017-2026-CGGC-{esp_sigla}-PEJGS.")

    ref_runs_text = ""
    for idx, ref_item in enumerate(refs_list):
        letra = chr(ord('a') + idx)
        prefix = ":  " if idx == 0 else "   "
        ref_runs_text += f"{prefix}{letra}) {_sanitize_xml_str(ref_item)}\n"
    
    p_rf_r.add_run(ref_runs_text.rstrip("\n"))

    # ----------------------------------------------------
    # 5. SALUDO Y CUERPO PRINCIPAL
    # ----------------------------------------------------
    p_saludo = doc.add_paragraph()
    p_saludo.paragraph_format.space_before = Pt(14)
    p_saludo.paragraph_format.space_after = Pt(8)
    p_saludo.add_run("De nuestra mayor consideración:")

    # Párrafo 1 (Exactamente redactado según la muestra de obra)
    p_body1 = doc.add_paragraph()
    p_body1.paragraph_format.line_spacing = 1.15
    p_body1.paragraph_format.space_before = Pt(0)
    p_body1.paragraph_format.space_after = Pt(8)
    p_body1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_body1.add_run(
        "Nos dirigimos a usted en atención a los documentos de las referencias b) y c), mediante "
        f"los cuales la Supervisión formuló observaciones a la propuesta técnica de {_sanitize_xml_str(esp_txt.lower())} "
        "presentada por el Contratista a través de los documentos de la referencia d)."
    )

    # Párrafo 2 (Presentación de informe técnico y absolución)
    p_body2 = doc.add_paragraph()
    p_body2.paragraph_format.line_spacing = 1.15
    p_body2.paragraph_format.space_before = Pt(0)
    p_body2.paragraph_format.space_after = Pt(8)
    p_body2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_body2.add_run(
        f"Al respecto, mediante la presente remitimos el Informe N.° 017-2026-CGGC-{esp_sigla}-PEJGS, "
        f"elaborado por el especialista en {_sanitize_xml_str(esp_txt.title())}, el cual contiene el análisis técnico, "
        "el sustento y la documentación actualizada para el levantamiento de las observaciones "
        "correspondientes a los siguientes elementos:"
    )

    # Lista con viñetas de ítems técnicos
    obs_txt = _sanitize_xml_str(payload.get("observacion") or "")
    items = []
    if obs_txt.strip():
        for line in obs_txt.splitlines():
            clean_l = _sanitize_xml_str(line).strip().lstrip("•-*0123456789.) ")
            if clean_l:
                items.append(clean_l)
    
    if not items:
        items = [
            "Ítem 1: Detalle y sustento técnico de la absolución formulada.",
            "Ítem 2: Especificaciones técnicas complementarias y replanteo de campo.",
            "Ítem 3: Planos de detalle y fichas técnicas actualizadas de los elementos observados."
        ]

    for item in items:
        p_bullet = doc.add_paragraph(style='List Bullet')
        p_bullet.paragraph_format.space_before = Pt(0)
        p_bullet.paragraph_format.space_after = Pt(3)
        p_bullet.paragraph_format.line_spacing = 1.15
        p_bullet.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r_b = p_bullet.add_run(_sanitize_xml_str(item))
        r_b.font.size = Pt(10)

    # Párrafo 3 (Conclusión formal)
    p_body3 = doc.add_paragraph()
    p_body3.paragraph_format.line_spacing = 1.15
    p_body3.paragraph_format.space_before = Pt(8)
    p_body3.paragraph_format.space_after = Pt(10)
    p_body3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_body3.add_run(
        "Por lo expuesto, se remite la presente comunicación para su conocimiento, evaluación y aprobación correspondiente."
    )

    # Despedida formal
    p_desp = doc.add_paragraph()
    p_desp.paragraph_format.space_before = Pt(4)
    p_desp.paragraph_format.space_after = Pt(48) # Espacio para la firma
    p_desp.add_run("Sin otro particular, quedamos de usted.")

    # ----------------------------------------------------
    # 6. PIE DE FIRMA OFICIAL
    # ----------------------------------------------------
    emisor = _sanitize_xml_str(payload.get("receptor") or payload.get("emisor") or "RO").upper()
    es_rl = "RL" in emisor or "LEGAL" in emisor
    cargo_firmante = "Representante Legal" if es_rl else "Residente de Obra"

    table_sign = doc.add_table(rows=1, cols=2)
    table_sign.alignment = WD_TABLE_ALIGNMENT.CENTER
    _eliminar_bordes_tabla(table_sign)
    table_sign.autofit = False
    table_sign.cell(0, 0).width = Inches(3.1)
    table_sign.cell(0, 1).width = Inches(3.37)

    # Firma oficial en columna derecha
    cell_sign = table_sign.cell(0, 1)
    p_sign = cell_sign.paragraphs[0]
    p_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sign.paragraph_format.space_before = Pt(0)
    p_sign.paragraph_format.space_after = Pt(2)
    p_sign.paragraph_format.line_spacing = 1.1

    r_line = p_sign.add_run("_________________________________________\n")
    r_line.font.bold = True
    r_emp1 = p_sign.add_run("CHINA GEZHOUBA GROUP COMPANY LIMITED\n")
    r_emp1.font.bold = True
    r_emp1.font.size = Pt(9.5)
    r_emp2 = p_sign.add_run("SUCURSAL PERÚ\n\n\n")
    r_emp2.font.bold = True
    r_emp2.font.size = Pt(9)
    r_ing = p_sign.add_run(f"Ing. [Nombre del {_sanitize_xml_str(cargo_firmante)}]\n")
    r_ing.font.bold = True
    r_ing.font.size = Pt(10)
    r_cip = p_sign.add_run(f"{_sanitize_xml_str(cargo_firmante.upper())} · CIP N° [XXXXXX]\n")
    r_cip.font.size = Pt(9)

    # Retornar buffer en memoria
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
